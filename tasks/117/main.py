from google.cloud import bigquery
from google.cloud import storage as st
from core.google import storage
from core.ports import Process
from io import BytesIO
from dateutil.relativedelta import relativedelta
from datetime import datetime as dt
import docx
from docx.shared import Pt



class Process117(Process):
    def __init__(self, **kwargs):
        self._output_bucket_name=kwargs["output_bucket_name"]
        self._blob_name = kwargs["blob_name"]
        self._exec_date=kwargs["exec_date"]
        self._query_date=kwargs["query_date"]
        self._header_date=kwargs["header_date"]
        self._bucket_name = kwargs["bucket_name"]
        self._project_name=kwargs["project_name"]
        self._dataset=kwargs["dataset"]
        self._codigo_reporte=kwargs["codigo_reporte"]
        self._codigo_dag=kwargs["codigo_dag"]
    def run(self):
        """the run method will execute instanciate the driver and call the other methods """
        self.run_query_and_save_to_gcs()

    def run_query_and_save_to_gcs(self):
        buffer=BytesIO()
        client = bigquery.Client()
        client_gcs=st.Client()
        
        query=storage.get_blob_as_string(f"gs://{self._bucket_name}/{self._codigo_dag}/queries/reporte-cmf-{self._codigo_reporte}.sql")
        parsed_query=self.parse_query(query)
        query_job = client.query(parsed_query)
        query_job.result()
        dataframe = query_job.to_dataframe()
        dataframe.to_parquet(buffer,compression='snappy')

        bucket = client_gcs.bucket(self._output_bucket_name)
        
        blob=bucket.blob(f"financial-reports/{self._codigo_reporte}/txt/{self._codigo_reporte}_{self._exec_date.replace('-','')}.txt")
        blob.upload_from_string(dataframe.to_csv(index=False,header=False),'text/plain')
        
        parquet_blob=bucket.blob(f"financial-reports/{self._codigo_reporte}/parquets/{self._codigo_reporte}_{self._exec_date.replace('-','')}.parquet")
        parquet_blob.upload_from_string(buffer.getvalue())
                

        query_2=storage.get_blob_as_string(f"gs://{self._bucket_name}/{self._codigo_dag}/queries/caratula-cmf-{self._codigo_reporte}.sql")
        parsed_query_2=self.parse_query(query_2)
        query_job_2 = client.query(parsed_query_2)
        query_job_2.result()
        dataframe_2 = query_job_2.to_dataframe()
        list_of_fields=sorted(
    [(i["orden"],i["Tipo"],i["Cuenta"]) for i in dataframe_2.drop(columns=["EXECUTION_DATE"]).to_dict(orient="records") ],
    key=lambda x : x[0]
    )
        caratula=self.create_caratula(list_of_fields)
        caratula.save(buffer)
        buffer.seek(1)
        caratula_blob=bucket.blob(f"financial-reports/{self._codigo_reporte}/caratulas/{self._codigo_reporte}_CARATULA_{self._exec_date.replace('-','')}.docx")
        caratula_blob.upload_from_string(buffer.read())
       
    def parse_query(self, query):
        return (
            query.replace(r"${project_name}", self._project_name)\
            .replace(r"${header_date}", self._header_date)\
            .replace(r"${query_date}", self._query_date)\
            .replace(r"${execution_date}", self._exec_date)
            .replace(r"${dataset}", self._dataset)\
            .replace(r"${codigo_reporte}", self._codigo_reporte.upper())
        )
    
    def create_caratula(self, list_of_fields):
        exec_date=self._header_date
        document= docx.Document()
        style=document.styles["Normal"]
        font=style.font
        font.name="Georgia"
        font.size=Pt(11)


        p=document.add_paragraph('Institución: Tenpo                                                                                    Código:730')
        run=p.add_run()
        run.add_break()
        p.add_run(f'Información correspondiente a la fecha: {exec_date}                               Archivo: {self._codigo_reporte}')

        p.style=document.styles["Normal"]

        table=document.add_table(rows=0,cols=2)
        table.style='Table Grid'
        table.autofit=False
        table.allow_autofit=False
        table.columns[0].width=4417200
        table.columns[1].width=1281600

        for element in list_of_fields:
            row_cells=table.add_row().cells
            row_cells[0].text=element[1]
            row_cells[1].text=str(element[2])
            run=row_cells[1].paragraphs[0].runs
            font=run[0].font
            font.name="Times New Roman"
            font.size=Pt(9)
        document.add_page_break()
        return document
